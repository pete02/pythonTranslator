from logging import root
import tkinter as tk
from tkinter import BooleanVar, Button, Canvas, Checkbutton, Frame, IntVar, filedialog, Text, OptionMenu,StringVar
import os
from logging import NullHandler
import time
from tkinter.constants import CENTER, FALSE, TRUE
import zipfile
import pyperclip
from selenium import webdriver
from selenium.webdriver.common.action_chains import ActionChains
from selenium.webdriver.firefox.options import Options
import docx
from docx2pdf import convert
import threading
import smtplib
from email.mime.base import MIMEBase 
from email.mime.multipart import MIMEMultipart 
from email.mime.text import MIMEText
from email import encoders
import docx2txt as d2t
import ctypes
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from googletrans import Translator
import shutil
import zipfile
import shutil
import xml.etree.ElementTree as ET
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.common.by import By
from concurrent.futures import ThreadPoolExecutor, as_completed

def send_msg(file):
    fromadd="petrik.rasanen@gmail.com"
    if(msgaddr.get()!="none"):
        toadd=msgaddr.get()
    else:
        toadd="petrik.rasanen@gmail.com"
    print("sending to "+toadd)
    msg= MIMEMultipart()
    msg['subject']="convert"
    msg['From']= fromadd
    msg['To']= toadd
    filename='C:\\Users\\petri\\Documents\\python\\books\\translated\\'+file+'.docx'
    attachment = open(filename, 'rb')

    part = MIMEBase("application", "octet-stream")

    part.set_payload(attachment.read())

    encoders.encode_base64(part)

    part.add_header("Content-Disposition",
    f"attachment; filename= "+file+".docx")

    msg.attach(part)

    msg = msg.as_string()


    server = smtplib.SMTP('smtp.gmail.com:587')
    server.ehlo()
    server.starttls()
    server.login(fromadd, '')
    server.sendmail(fromadd, toadd, msg)
    server.quit()

def addFile():
    global files
    for widget in frame.winfo_children():
        widget.destroy()
    filename = filedialog.askopenfilename(initialdir=r"C:\Users\petri\Documents\python\books\japanese", title="Select File",
    filetypes=(("docx", "*.docx"), ("all files", "*.*")))
    print(filename)
    files.append(filename)
    for file in files:
        label = tk.Label(frame, text=file, bg="#D3D3D3")
        label.pack()

def rmFi():
    global files
    for widget in frame.winfo_children():
        widget.destroy()
    files = []

def extract(file,im):
    d2t.process(file,im)
    

def getName(ele):
    root=ET.fromstring(ele.xml)
    j=root[0][0][0].attrib
    return(str(j["name"]).replace(".jpg",".jpeg"))

def getNames(doc):
    l=[]
    for elem in doc.element.getiterator():
        if('graphicData'in elem.tag):
            l.append(getName(elem))
    return(l)


def bulk(file,img):
    doc=docx.Document(file)
    w=""
    l=[]
    imnames=getNames(doc)
    extract(file,img)
    i=0
    p=0
    for paragraph in doc.paragraphs:
        i=i+1
        if 'graphicData' in paragraph._p.xml:
            print("img:",imnames[p])
            l.append(["img",imnames[p]])
            p=p+1
        else:
            text = paragraph.text
            if(text != ""):
                text=text.replace('「','"') 
                text=text.replace('」','"')
                text=text.replace('『','"') 
                text=text.replace('』','"')
                if(len(w)+len(text)<1000):
                    w=w+text+"\n"
                else:
                    l.append(["text",w])
                    w=text + "\n"

    l.append(w)
    return l


def set_Diver():
    firefox_profile = webdriver.FirefoxProfile()
    firefox_profile.set_preference("browser.privatebrowsing.autostart", True)
    options = Options()
    options.headless = True
    driver = webdriver.Firefox(executable_path="C:\\Users\\petri\\Documents\\python\\webdriver\\geckodriver.exe",firefox_profile=firefox_profile,options=options )
    deepl_url = 'https://www.deepl.com/en/translator#ja/en/'
    driver.get(deepl_url)
    print("run driver")
    return driver


def check(text):
    if(Translator().detect(text)!='ja'):
        return text
    else:
        return None


def translate():
    try:
        shutil.rmtree("C:\\Users\\petri\\Documents\\python\\imgs")
    except:
        pass
    os.mkdir("C:\\Users\\petri\\Documents\\python\\imgs")
    ctypes.windll.kernel32.SetThreadExecutionState(0x80000002)
    name=""
    print("tenry:"+ etry.get())
    doc2 = docx.Document()
    #print("files:", files)
    if(files != []):
        
        for file in files:
            driver=set_Diver()
            copiedtext=""
            textlist=bulk(file,"C:\\Users\\petri\\Documents\\python\\imgs")
            lis=os.listdir("C:\\Users\\petri\\Documents\\python\\imgs")
            k=len(textlist)
            print(k)
            root.update_idletasks()
            i=0
            h=0
            all=0
            googl=0
            prevcontent = "salfkhadsgouhözxv"
            prevtext = "kajgfaoölighasdöogv"
            content=""
            # Send the text
            aika=True
            prevc=""
            for textl in textlist:
                all=+1
                same=True
                while(same):      
                    if(textl[0]=="text"):
                        start_time = time.time()
                        text=textl[1]
                        time.sleep(1)
                        content=google(text)
                        print("c",content,"e")
                        if(prevcontent !=content or prevtext == text):
                            c=content.split(" ")
                            with ThreadPoolExecutor() as executor:
                                content=executor.map(check,c)
                            content=" ".join(content)
                            print("cont:",content,"end")
                            same=False
                            content=content.replace("Keibobo","Kiri-boy" )
                            content=content.replace("Keebobo","Kiri-boy" )
                            content=content.replace("Kiibobo","Kiri-boy" )
                            content=content.replace("key boy","Kiri-boy" )
                            content=content.replace("Kee-bo-bo","Kiri-boy" )
                            content=content.replace("Kiku Kikuoka","Kikuoka" )
                            content=content.replace("<built-in function input>","")
                            print("saving")
                            h=h+1
                            j=(k-h)
                            if(aika==True): 
                                ti=time.time()-start_time
                                aika=False
                            titext=str(round((ti*j)/60,2))
                            t="estimated time: " + titext + "min"
                            print(t)
                            for widget in frame2.winfo_children():
                                widget.destroy()
                            label = tk.Label(frame2, text=t,bg="#D3D3D3")
                            label.pack()
                            root.update_idletasks()
                            paragraph = doc2.add_paragraph(content)
                            prevcontent=content
                            prevtext=text
                        else:
                            driver.quit()
                            driver=set_Diver()
                            print("same text")
                        try:
                            driver.manage().deleteCookie("deepl.com")
                            
                        except:
                            pass
                        try:
                            driver.manage().deleteCookie("www.deepl.com")
                            
                        except:
                            pass
                        content=""
                    else:
                        same=False
                        if(textl[0]=="img"):
                            doc2.add_picture("C:\\Users\\petri\\Documents\\python\\imgs\\"+textl[1],Cm(15.5))
                            last_paragraph = doc2.paragraphs[0] 
                            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            if(etry.get()!=NullHandler and etry.get()!=""):
                name=etry.get()
                saveD='C:\\Users\\petri\\Documents\\python\\books\\translated\\'+etry.get()+'.docx'
            else: 
                if(textl[0]=="img"):
                    saveD='C:\\Users\\petri\\Documents\\python\\translated.docx'
            i=i+1
            doc2.save(saveD)
            print("used google",googl/all)
            if(sendMS.get()):
                print("starting to send")
                send_msg(name)
            else:
                print("not sending")
            shutil.rmtree("C:\\Users\\petri\\Documents\\python\\imgs")
        driver.quit()
        ctypes.windll.kernel32.SetThreadExecutionState(0x80000000)


def google(text):
    trans=Translator()
    ja=True
    i=0
    while(ja):
        i+=1
        out=trans.translate(text)
        if(trans.detect(out.text)!='ja'):
            ja=False
        if(i>10):
            print("looped ten times")
            time.sleep(10)
    return out.text

def do():
    if(len(files)>0):
        threading.Thread(target=translate).start()

root = tk.Tk()
root.configure(height=700, width=700,bg="#263D42")
files = []


root.geometry("500x500")
#canvas = tk.Canvas(root, height=500, width=700, bg= "#263D42")
#canvas.pack()

frame = tk.Frame(root, bg= "#D3D3D3")
frame.place(relheight=0.5, relwidth= 0.8, relx=0.1, rely=0.1)

frame2 = tk.Frame(root, bg= "#D3D3D3")
frame2.place(relheight=0.05, relwidth= 0.4, relx=0.3, rely=0.025)

frame3 = tk.Frame(root, bg= "#D3D3D3")
frame3.place(height=20, width= 200, relx=0.3, rely=0.61)

openFile = tk.Button(root, text="Open File", padx=10, pady=5, fg="#D3D3D3", bg="#263D42", command= addFile)
openFile.place(relx=0.5, anchor="c", rely=0.7)


run= tk.Button(root, text="Run Translation", padx=10, pady=5, fg="#D3D3D3", bg="#263D42", command= do)
run.place(relx=0.5, anchor="c", rely=0.84)

dest= tk.Button(root, text="Remove Files", padx=10, pady=5, fg="#D3D3D3", bg="#263D42", command= rmFi)
dest.place(relx=0.5, anchor="c", rely=0.77)


sendMS=BooleanVar()
check=Checkbutton(root,text="send file",variable=sendMS, fg="#D3D3D3", bg="#263D42",onvalue=TRUE, offvalue=FALSE)
check.place(relx=0.1, anchor="c", rely=0.77)

OPTIONS = [
"none",
"petrik.rasanen@gmail.com",
"annika.ahlholm@kindle.com"
]


msgaddr = StringVar(root)
msgaddr.set(OPTIONS[0]) # default value

w = OptionMenu(root, msgaddr, *OPTIONS).place(relx=0.20,anchor= "c", rely=0.70)


etry=tk.Entry(frame3, fg="#263D42",width=40,bg='#D3D3D3')
etry.pack()
root.mainloop()

