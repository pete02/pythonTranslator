from logging import root
import tkinter as tk
from tkinter import BooleanVar, Button, Canvas, Checkbutton, Frame, IntVar, filedialog, Text, OptionMenu,StringVar
import os
from logging import NullHandler
import time
from tkinter.constants import CENTER, FALSE, TRUE
import docx
import xml.etree.ElementTree as ET
import threading
import smtplib
from email.mime.base import MIMEBase 
from email.mime.multipart import MIMEMultipart 
from email import encoders
import docx2txt as d2t
import ctypes
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import shutil
import concurrent.futures
from easynmt import EasyNMT


model = EasyNMT('opus-mt')




def send_msg(file):
    password= None
    with open('secret.txt') as f:
        password = f.readlines()[0]

    fromadd="petrik.rasanen@gmail.com"
    if(msgaddr.get()!="none"):
        toadd=msgaddr.get()
    else:
        toadd="petrik.rasanen@gmail.com"
    print("sending to "+toadd)
    msg= MIMEMultipart()
    msg['subject']="convert"
    msg['From']= fromadd
    msg['To']= toadd
    filename='C:\\Users\\petri\\Documents\\python\\books\\translated\\'+file+'.docx'
    attachment = open(filename, 'rb')

    part = MIMEBase("application", "octet-stream")

    part.set_payload(attachment.read())

    encoders.encode_base64(part)

    part.add_header("Content-Disposition",
    f"attachment; filename= "+file+".docx")

    msg.attach(part)

    msg = msg.as_string()


    server = smtplib.SMTP('smtp.gmail.com:587')
    server.ehlo()
    server.starttls()
    server.login(fromadd, password)
    server.sendmail(fromadd, toadd, msg)
    server.quit()

def addFile():
    global files
    for widget in frame.winfo_children():
        widget.destroy()
    filename = filedialog.askopenfilename(initialdir=r"C:\Users\petri\Documents\python\books\japanese", title="Select File",
    filetypes=(("docx", "*.docx"), ("all files", "*.*")))
    print(filename)
    files.append(filename)
    for file in files:
        label = tk.Label(frame, text=file, bg="#D3D3D3")
        label.pack()

def rmFi():
    global files
    for widget in frame.winfo_children():
        widget.destroy()
    files = []

def extract(p,im):
    text=d2t.process(p,im)


def getName(ele):
    root=ET.fromstring(ele.xml)
    j=root[0][0][0].attrib
    return(str(j["name"]).replace("jpg","jpeg"))

def getNames(doc):
    l=[]
    for elem in doc.element.getiterator():
        if('graphicData'in elem.tag):
            l.append(getName(elem))
    return(l)


def bulk(file,img):
    doc=docx.Document(file)
    w=""
    l=[]
    imnames=getNames(doc)
    extract(file,img)
    i=0
    p=0
    for paragraph in doc.paragraphs:
        i=i+1
        if 'graphicData' in paragraph._p.xml:
            print("img:",imnames[p])
            l.append(["img",imnames[p],i])
            p=p+1
        else:
            text = paragraph.text
            if(text != ""):
                if(len(w)+len(text)<1000):
                    w=w+text+"\n"
                else:
                    l.append(["text",w,i])
                    w=text + "\n"

    l.append(w)
    return l


def tran2(text):
    return[text[0],model.translate(text[1], target_lang='en'),text[2]]

def translate():
    try:
        shutil.rmtree("C:\\Users\\petri\\Documents\\python\\imgs")
    except:
        pass
    os.mkdir("C:\\Users\\petri\\Documents\\python\\imgs")
    ctypes.windll.kernel32.SetThreadExecutionState(0x80000002)
    name=""
    print("tenry:"+ etry.get())
    doc2 = docx.Document()
    #print("files:", files)
    if(files != []):
        i=1
        for file in files:
            start = time.perf_counter()
            print("start translaitting")
            list=bulk(file,"C:\\Users\\petri\\Documents\\python\\imgs")
            lis=os.listdir("C:\\Users\\petri\\Documents\\python\\imgs")
            txtl=[]
            imgl=[]
            txtl2=[]
            root.update_idletasks()

            for l in list:
                if(l[0]=="text"):
                    txtl.append(l)
                if(l[0]=="img"):
                    imgl.append(l)
            with concurrent.futures.ThreadPoolExecutor(max_workers=5) as excecutor:
                results=excecutor.map(tran2,txtl)
                for result in results:
                    txtl2.append(result)
            list=imgl+txtl2
            list=sorted(list,key=lambda x:x[2])
            for l in list:
                print(l[2])
                if(l[0]=="img"):
                    
                    doc2.add_picture("C:\\Users\\petri\\Documents\\python\\imgs\\"+l[1],Cm(15.5))
                    last_paragraph = doc2.paragraphs[0] 
                    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    os.remove("C:\\Users\\petri\\Documents\\python\\imgs\\"+l[1])
                if(l[0]=="text"):
                    t=l[1]
                    doc2.add_paragraph(t)
            if(etry.get()!=NullHandler and etry.get()!=""):
                saveD='C:\\Users\\petri\\Documents\\python\\books\\translated\\'+etry.get()+'.docx'
            else: 
                saveD='C:\\Users\\petri\\Documents\\python\\translated.docx'
            doc2.save(saveD)
            if(sendMS.get()):
                print("starting to send")
                send_msg(etry.get())
            else:
                print("not sending")
            os.rmdir("C:\\Users\\petri\\Documents\\python\\imgs")
            print("done transalting")
            print("time to finish:",time.perf_counter()-start)
        ctypes.windll.kernel32.SetThreadExecutionState(0x80000000)


def do():
    if(len(files)>0):
        threading.Thread(target=translate).start()

root = tk.Tk()
root.configure(height=700, width=700,bg="#263D42")
files = []


root.geometry("500x500")
#canvas = tk.Canvas(root, height=500, width=700, bg= "#263D42")
#canvas.pack()

frame = tk.Frame(root, bg= "#D3D3D3")
frame.place(relheight=0.5, relwidth= 0.8, relx=0.1, rely=0.1)

frame2 = tk.Frame(root, bg= "#D3D3D3")
frame2.place(relheight=0.05, relwidth= 0.4, relx=0.3, rely=0.025)

frame3 = tk.Frame(root, bg= "#D3D3D3")
frame3.place(height=20, width= 200, relx=0.3, rely=0.61)

openFile = tk.Button(root, text="Open File", padx=10, pady=5, fg="#D3D3D3", bg="#263D42", command= addFile)
openFile.place(relx=0.5, anchor="c", rely=0.7)


run= tk.Button(root, text="Run Translation", padx=10, pady=5, fg="#D3D3D3", bg="#263D42", command= do)
run.place(relx=0.5, anchor="c", rely=0.84)

dest= tk.Button(root, text="Remove Files", padx=10, pady=5, fg="#D3D3D3", bg="#263D42", command= rmFi)
dest.place(relx=0.5, anchor="c", rely=0.77)


sendMS=BooleanVar()
check=Checkbutton(root,text="send file",variable=sendMS, fg="#D3D3D3", bg="#263D42",onvalue=TRUE, offvalue=FALSE)
check.place(relx=0.1, anchor="c", rely=0.77)

OPTIONS = [
"none",
"petrik.rasanen@gmail.com",
"annika.ahlholm@kindle.com"
]


msgaddr = StringVar(root)
msgaddr.set(OPTIONS[0]) # default value

w = OptionMenu(root, msgaddr, *OPTIONS).place(relx=0.20,anchor= "c", rely=0.70)


etry=tk.Entry(frame3, fg="#263D42",width=40,bg='#D3D3D3')
etry.pack()
root.mainloop()


